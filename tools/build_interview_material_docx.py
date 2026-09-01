"""Convert the maintained interview-material Markdown into a print-ready DOCX."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "outputs" / "应用+微调方面实习_项目简历及面试材料.md"
OUTPUT = ROOT / "outputs" / "Python应用实习_项目简历与面试材料.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "666666"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
USABLE_WIDTH_DXA = 9360


def set_cjk_font(run, ascii_font: str = "Calibri", east_asia_font: str = "Microsoft YaHei"):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)
    fonts.set(qn("w:eastAsia"), east_asia_font)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_paragraph_border(paragraph, side: str, color: str, size: int = 8, space: int = 6):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_cjk_font(run)


def set_table_geometry(table, widths: list[int]):
    """Apply fixed DXA geometry to table width, grid, indent and cells."""
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(USABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid_cols = tbl.tblGrid.gridCol_lst
    for index, width in enumerate(widths):
        grid_cols[index].set(qn("w:w"), str(width))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    quote = styles["Quote"]
    quote.font.name = "Calibri"
    quote.font.size = Pt(10.5)
    quote.font.color.rgb = RGBColor.from_string(MUTED)
    quote.paragraph_format.left_indent = Inches(0.24)
    quote.paragraph_format.right_indent = Inches(0.08)
    quote.paragraph_format.space_before = Pt(3)
    quote.paragraph_format.space_after = Pt(5)
    quote.paragraph_format.line_spacing = 1.2

    if "Reference Answer" not in styles:
        reference = styles.add_style("Reference Answer", WD_STYLE_TYPE.PARAGRAPH)
    else:
        reference = styles["Reference Answer"]
    reference.base_style = normal
    reference.font.name = "Calibri"
    reference.font.size = Pt(11)
    reference.paragraph_format.space_before = Pt(0)
    reference.paragraph_format.space_after = Pt(6)
    reference.paragraph_format.line_spacing = 1.25


def add_inline(paragraph, text: str):
    """Render minimal Markdown inline emphasis without changing the source wording."""
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_cjk_font(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            set_cjk_font(run)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            set_cjk_font(run, "Consolas")
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            set_cjk_font(run)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_cjk_font(run)


def add_regular_paragraph(doc: Document, text: str, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    add_inline(paragraph, text.rstrip())
    return paragraph


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_markdown_table(doc: Document, rows: list[list[str]]):
    columns = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    if columns == 3:
        widths = [2200, 3400, 3760]
    elif columns == 2:
        widths = [2700, 6660]
    else:
        base = USABLE_WIDTH_DXA // columns
        widths = [base] * columns
        widths[-1] += USABLE_WIDTH_DXA - sum(widths)
    set_table_geometry(table, widths)
    for row_index, source_row in enumerate(rows):
        for column_index, value in enumerate(source_row):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.15
            add_inline(paragraph, value)
            for run in paragraph.runs:
                run.font.size = Pt(9.5)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for run in paragraph.runs:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_heading(doc: Document, level: int, text: str, is_first_title: bool):
    if level == 1 and is_first_title:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(20)
        paragraph.paragraph_format.space_after = Pt(16)
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        set_cjk_font(run)
        return paragraph
    style = "Heading 1" if level == 2 else "Heading 2" if level == 3 else "Heading 3"
    paragraph = doc.add_paragraph(style=style)
    add_inline(paragraph, text)
    if level == 2:
        paragraph.paragraph_format.page_break_before = True
    return paragraph


def build_document(source: Path, output: Path):
    if not source.exists():
        raise FileNotFoundError(source)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    doc.core_properties.title = "Python 应用实习｜项目简历与面试材料"
    doc.core_properties.subject = "项目简历、自我介绍与模拟面试问答"
    doc.core_properties.author = ""
    doc.core_properties.comments = "由 Markdown 面试材料转换"

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(2)
    header_run = header.add_run("PYTHON 应用实习｜项目简历与面试材料")
    header_run.font.size = Pt(8.5)
    header_run.font.color.rgb = RGBColor.from_string(MUTED)
    set_cjk_font(header_run)
    set_paragraph_border(header, "bottom", "D7DBE2", size=6, space=4)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("第 ")
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = RGBColor.from_string(MUTED)
    set_cjk_font(footer_run)
    add_field(footer, "PAGE")
    footer_run = footer.add_run(" 页")
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = RGBColor.from_string(MUTED)
    set_cjk_font(footer_run)

    lines = source.read_text(encoding="utf-8").splitlines()
    first_title = True
    index = 0
    bullet_re = re.compile(r"^\s*[-*]\s+(.+)$")
    number_re = re.compile(r"^\s*\d+\.\s+(.+)$")
    heading_re = re.compile(r"^(#{1,3})\s+(.+)$")
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue
        heading = heading_re.match(stripped)
        if heading:
            add_heading(doc, len(heading.group(1)), heading.group(2), first_title)
            if len(heading.group(1)) == 1:
                first_title = False
            index += 1
            continue
        if stripped == "---":
            index += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                candidate = lines[index].strip()
                if not re.fullmatch(r"\|?\s*[:-]+(?:\s*\|\s*[:-]+)+\s*\|?", candidate):
                    table_lines.append(parse_table_row(candidate))
                index += 1
            if table_lines:
                add_markdown_table(doc, table_lines)
            continue
        if stripped.startswith("> "):
            paragraph = add_regular_paragraph(doc, stripped[2:].rstrip(), style="Quote")
            set_paragraph_border(paragraph, "left", BLUE, size=18, space=8)
            index += 1
            continue
        bullet = bullet_re.match(raw)
        if bullet:
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet.group(1).rstrip())
            index += 1
            continue
        numbered = number_re.match(raw)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            add_inline(paragraph, numbered.group(1).rstrip())
            index += 1
            continue
        style = "Reference Answer" if stripped.startswith("**参考回答：**") else None
        add_regular_paragraph(doc, stripped, style=style)
        index += 1

    # The document is intentionally kept as one portrait section for consistent headers and footers.
    assert len(doc.sections) == 1
    doc.save(output)


if __name__ == "__main__":
    source_arg = Path(sys.argv[1]) if len(sys.argv) > 1 else SOURCE
    output_arg = Path(sys.argv[2]) if len(sys.argv) > 2 else OUTPUT
    build_document(source_arg, output_arg)
    print(output_arg)
