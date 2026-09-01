from pathlib import Path
import os
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "测试实习生笔试题_V1_完整作答.md"
OUTPUT = ROOT / os.environ.get("DOCX_OUTPUT_NAME", "测试实习生笔试题_V1_完整作答.docx")


def set_run_font(run, name="微软雅黑", size=10.5, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, header=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text.strip())
    set_run_font(run, size=8.5 if not header else 8.8, bold=header)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if header:
        shade_cell(cell, "D9EAF7")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, rows):
    if not rows:
        return
    parsed = [[part.strip() for part in row.strip().strip("|").split("|")] for row in rows]
    columns = max(len(row) for row in parsed)
    table = doc.add_table(rows=len(parsed), cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(parsed):
        for c_idx in range(columns):
            value = row[c_idx] if c_idx < len(row) else ""
            set_cell_text(table.cell(r_idx, c_idx), value, header=r_idx == 0)
    repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph(style="No Spacing")
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.right_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=8.5, color=(45, 45, 45))
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)


def add_rich_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(5)
    # Basic Markdown emphasis/code handling while preserving readable text.
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=9.5, color=(40, 80, 110))
        elif part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = p.add_run(part)
            set_run_font(run)
    return p


def configure_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 20, (31, 78, 121)),
        ("Heading 1", 15, (31, 78, 121)),
        ("Heading 2", 12, (46, 116, 181)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(*color)


def convert():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("```"):
            code = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i])
                i += 1
            add_code_block(doc, code)
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].lstrip().startswith("|---"):
            table_rows = [line]
            i += 1
            while i < len(lines) and lines[i].startswith("|"):
                if not re.match(r"^\|\s*:?-+", lines[i]):
                    table_rows.append(lines[i])
                i += 1
            add_table(doc, table_rows)
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].strip())
            set_run_font(run, size=20, bold=True, color=(31, 78, 121))
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            run = p.add_run(line[3:].strip())
            set_run_font(run, size=15, bold=True, color=(31, 78, 121))
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            run = p.add_run(line[4:].strip())
            set_run_font(run, size=12, bold=True, color=(46, 116, 181))
        elif line.startswith("> "):
            p = add_rich_paragraph(doc, line[2:].strip())
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.right_indent = Cm(0.6)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(90, 90, 90)
        elif line == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            border = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "B7C9D6")
            border.append(bottom)
            p._p.get_or_add_pPr().append(border)
        elif re.match(r"^\d+\.\s", line) or line.startswith("- "):
            content = re.sub(r"^(?:\d+\.\s|- )", "", line).strip()
            p = add_rich_paragraph(doc, content, style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.6)
        else:
            add_rich_paragraph(doc, line)
        i += 1

    # Add a simple footer with page numbering field.
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("测试实习生笔试题 V1 完整作答  |  第 ")
    set_run_font(run, size=8, color=(120, 120, 120))
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)
    doc.save(OUTPUT)


if __name__ == "__main__":
    convert()
    print(OUTPUT)
